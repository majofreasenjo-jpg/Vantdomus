import sys
import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

def build_header(doc, recipient_name, recipient_role, reference, subject, contract_id):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Emisión Autogenerada, {datetime.now().strftime('%d de %B de %Y')}")
    
    p = doc.add_paragraph()
    p.add_run("\nSEÑORES\n").bold = True
    p.add_run(f"{recipient_name.upper()}\n").bold = True
    p.add_run(f"Atte. {recipient_role}\n")
    p.add_run("Presente.\n")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"REF: {reference.upper()}\n").bold = True
    p.add_run(f"MATERIA: {subject.upper()}\n").bold = True
    p.add_run(f"CONTRATO: {contract_id}").bold = True
    
    doc.add_paragraph("De nuestra consideración:")
    
def add_justified(doc, text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def main():
    if len(sys.argv) < 3:
        print("Uso: python letter_generator.py <params.json> <out_dir>")
        sys.exit(1)
        
    params_file = sys.argv[1]
    out_dir = sys.argv[2]
    
    with open(params_file, "r", encoding="utf-8") as f:
        args = json.load(f)
        
    contract_id = args.get("contract_id", "CT-GENERIC")
    recipient_name = args.get("recipient_name", "COMPAÑÍA CLIENTE")
    recipient_role = args.get("recipient_role", "Administración de Contratos")
    level = args.get("severity_level", "amistoso").lower()
    subject = args.get("subject", "Notificación Comercial")
    amount = args.get("amount", "")
    facts = args.get("facts_description", "Detalle de los hechos sucedidos...")
    legal_arguments = args.get("legal_arguments", "")
    
    # Configure tones based on severity
    if level == "amistoso":
        reference = "SOLICITUD DE REVISIÓN Y SOLUCIÓN AMISTOSA"
        opening = f"A través de la presente, nuestra empresa se dirige cordialmente a ustedes con el propósito de poner en su conocimiento la siguiente situación material y solicitar una evaluación comercial conjunta."
        closing_demand = "Quedamos a su entera disposición para agendar una reunión técnica y revisar estos antecedentes con miras a la mayor colaboración y equidad posible frente al cierre."
    elif level == "reclamo":
        reference = "PRESENTACIÓN DE RECLAMO CONTRACTUAL Y NUEVOS ANTECEDENTES"
        opening = f"Mediante el presente documento, venimos en interponer formalmente a su consideración un RECLAMO CONTRACTUAL al amparo de las Condiciones Generales del Contrato, respecto a los hechos que a continuación se exponen."
        closing_demand = "Notificamos la reserva expresa de nuestros Mayores Gastos Generales y exigimos la activación formal de la Mesa de Solución Amistosa estipulada en el contrato para resarcir los desvíos reportados."
    elif level == "arbitral":
        reference = "NOTIFICACIÓN FORMAL DE EXISTENCIA DE CONTROVERSIA PRE-ARBITRAL E INCUMPLIMIENTO"
        opening = f"Por medio de la presente, emplazamos y notificamos formalmente la existencia de controversia jurídica y económica grave respecto de la ejecución operativa de nuestro contrato, imputable a incumplimientos principales del mandante."
        closing_demand = "Téngase la presente como intimación formal. Exigimos un período de avenimiento y citación de la Gerencia General en un plazo no superior a cinco (5) días hábiles, so pena de trasladar la cuantificación de Daño Emergente, Costos Financieros y Lucro Cesante a Tribunales Arbitrales y demandar la exceptio non adimpleti contractus por vuestros reiterados incumplimientos."
    else:
        reference = "COMUNICACIÓN FORMAL"
        opening = "Nos dirigimos a ustedes para formalizar los siguientes antecedentes:"
        closing_demand = "Quedamos a la espera de su respuesta."
        
    if amount:
        subject += f" (Monto Involucrado: {amount})"
        
    doc = Document()
    set_style(doc)
    build_header(doc, recipient_name, recipient_role, reference, subject, contract_id)
    
    add_justified(doc, "1. ANTECEDENTES Y OBJETO", bold=True)
    add_justified(doc, opening)
    
    add_justified(doc, "2. EXPOSICIÓN DE LOS HECHOS", bold=True)
    add_justified(doc, facts)
    
    if legal_arguments:
        add_justified(doc, "3. FUNDAMENTOS CONTRACTUALES Y DAÑOS", bold=True)
        add_justified(doc, legal_arguments)
    
    add_justified(doc, f"{'4' if legal_arguments else '3'}. EXIGENCIA Y EMPLAZAMIENTO", bold=True)
    add_justified(doc, closing_demand)
    
    add_justified(doc, "Se despide atentamente,\n\n\n\n________________________________\nFirma Autorizada\nContratista")
    
    filename = f"Carta_{level.capitalize()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    # Safe filename
    for c in ["/", "\\", "?", "%", "*", ":", "|", '"', "<", ">"]:
        filename = filename.replace(c, "")
        
    out_path = os.path.join(out_dir, filename)
    doc.save(out_path)
    
    # Print exactly the output path for the caller to parse
    print(out_path)

if __name__ == "__main__":
    main()
