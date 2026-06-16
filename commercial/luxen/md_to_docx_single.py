import os
import subprocess
import sys

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

def md_to_docx(md_path, docx_path):
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:].replace('**', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].replace('**', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].replace('**', ''), level=3)
        elif line.startswith('- ') or line.startswith('* '):
            clean_line = line[2:].replace('**', '')
            doc.add_paragraph(clean_line, style='List Bullet')
        elif line.startswith('> [!'):
            pass # skip alerts specific syntax
        elif line.startswith('> '):
            doc.add_paragraph(line[2:].replace('**', ''), style='List Paragraph')
        elif line == '---':
            pass # skip horizontal lines
        else:
            # Paragraph
            doc.add_paragraph(line.replace('**', '').replace('_', ''))
            
    doc.save(docx_path)
    print(f"Saved: {docx_path}")

md_files = ["resumen_modelo_tarifas.md"]
dr = r"d:\Aplicaciones de Juegos\VantDomus\Luxen_Comercial"

for md in md_files:
    path_in = os.path.join(dr, md)
    path_out = os.path.join(dr, md.replace('.md', '.docx'))
    if os.path.exists(path_in):
        md_to_docx(path_in, path_out)
        os.remove(path_in)
        print(f"Borrando origen: {path_in}")
